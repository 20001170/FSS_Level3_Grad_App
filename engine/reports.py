"""
Report builders. Produce the three downloadable deliverables from a summary:
  build_xlsx(results, path)         -> full review spreadsheet (all students)
  build_docx(summary, path)         -> formatted Word report
  build_pdf(summary, path)          -> PDF (rendered from the docx via LibreOffice if
                                       available, else a direct ReportLab PDF)
"""
import os
import shutil
import subprocess

from .settings import bands_description, all_band_eras, CREDIT_FLOORS


NAVY = "1F4E78"


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------
def build_xlsx(results, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Eligibility Review"

    title_font = Font(bold=True, size=13, name="Arial")
    sub_font = Font(italic=True, size=9, name="Arial")
    warn_font = Font(italic=True, size=9, name="Arial", color="C00000")
    header_font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    header_fill = PatternFill("solid", start_color=NAVY)
    normal = Font(name="Arial", size=10)
    green = PatternFill("solid", start_color="C6EFCE")
    yellow = PatternFill("solid", start_color="FFEB9C")
    grey = PatternFill("solid", start_color="E7E6E6")
    amber = PatternFill("solid", start_color="FFF2CC")
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    n = len(results)
    ws["A1"] = f"Level 3 Graduation Eligibility Review (n = {n})"
    ws["A1"].font = title_font
    ws["A2"] = ("Each student checked against the required-course list from the handbook edition matching "
                "their entry year. Course codes (not titles) are the basis of the check.")
    ws["A2"].font = sub_font
    ws["A3"] = ("Records are split per student (two students may share a physical page). Credit floors include "
                "EC exemptions. Class of degree uses the WEIGHTED GPA (Level I weighted zero; Levels II & III, "
                "passed courses).")
    ws["A3"].font = sub_font
    _band_eras = "  |  ".join(f"{lab}: {desc}" for lab, desc in all_band_eras())
    ws["A4"] = (f"Class-of-degree bands (by entry year) — {_band_eras}. "
                f"Credit floors: {CREDIT_FLOORS['total']} total / {CREDIT_FLOORS['level1']} Level I / "
                f"{CREDIT_FLOORS['level23']} Levels II-III. "
                "Each student is classified on the bands in force in their year of entry. "
                "Where weighted and printed GPA give different classes the weighted-GPA cell is highlighted "
                "(confirm with registry). This tool does not decide elective-slot counts, the Management 30/30 "
                "split, EX-exemption validity, rate-of-progress, or BOE annotations. Final sign-off: examiner.")
    ws["A4"].font = warn_font

    headers = ["Name", "ID", "Page", "Programme", "Entry Year", "Era",
               "Printed Degree GPA", "Weighted GPA (L2/3)",
               "Total Credits (incl. EC)", "EC Exempt Credits",
               "L1 Earned (incl. EC)", "L23 Earned",
               "Named-Course Check", "Missing Level I", "Missing Level II/III",
               "Class of Degree (weighted)", "Verdict"]
    start = 6
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=start, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = border

    def sort_key(r):
        v = r.get("verdict", "")
        if v.startswith("ELIGIBLE"):
            return (0, -(r.get("weighted_gpa") or 0), r["name"])
        if v.startswith("CREDIT FLOORS MET"):
            return (1, 0, r["name"])
        miss = len(r.get("missing_level1", [])) + len(r.get("missing_level23_named", []))
        return (2, miss, r["name"])

    row = start + 1
    for r in sorted(results, key=sort_key):
        checked = r.get("checked")
        vname = "Checked" if checked is True else ("Credit-only" if checked == "credit_only" else "NOT CHECKED")
        wg = r.get("weighted_gpa")
        vals = [
            r["name"], r["id"], r.get("page", "—"), r["programme"], r["entry_year"],
            r.get("era", "—") or "—", r["gpa"], (wg if wg is not None else "—"),
            r.get("total_credits_effective", r.get("total_credits")),
            r.get("ec_total", 0),
            r.get("level1_earned", "—"), r.get("level23_earned", "—"),
            vname,
            ", ".join(r.get("missing_level1", [])) or "—",
            ", ".join(r.get("missing_level23_named", [])) or "—",
            r.get("class_of_degree", "") or "—",
            r.get("verdict", ""),
        ]
        for c, v in enumerate(vals, start=1):
            cell = ws.cell(row=row, column=c, value=v)
            cell.font = normal
            cell.border = border
            cell.alignment = Alignment(wrap_text=True, vertical="top") if c == 17 else Alignment(vertical="top")
        v = r.get("verdict", "")
        if v.startswith("ELIGIBLE"):
            for c in range(1, 18):
                ws.cell(row=row, column=c).fill = green
            ws.cell(row=row, column=16).font = Font(name="Arial", size=10, bold=True)
            if r.get("class_differs"):
                ws.cell(row=row, column=8).fill = amber
        elif v.startswith("NOT ELIGIBLE — numeric OK"):
            for c in range(1, 18):
                ws.cell(row=row, column=c).fill = yellow
        elif checked == "credit_only":
            for c in range(1, 18):
                ws.cell(row=row, column=c).fill = grey
        row += 1

    widths = [26, 11, 6, 21, 10, 14, 11, 12, 13, 11, 13, 11, 13, 24, 28, 25, 48]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A7"
    ws.row_dimensions[4].height = 50

    wb.save(path)
    return path


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def build_docx(summary, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    navy = RGBColor(0x1F, 0x4E, 0x78)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), hexcolor)
        tcPr.append(sh)

    def heading(text, size=16):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = navy
        return p

    title = doc.add_paragraph()
    r = title.add_run("Level 3 Graduation Eligibility Review")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = navy

    c = summary["counts"]
    sub = doc.add_paragraph()
    sub.add_run(f"Faculty of Social Sciences — {c['total']} Level 3 records assessed").font.size = Pt(11)
    note = doc.add_paragraph()
    nr = note.add_run("Each student assessed against the required-course list from the handbook edition "
                      "matching their year of entry. Course codes (not titles) are the basis of the check. "
                      "The 'Page' column locates each transcript in the academic record.")
    nr.italic = True
    nr.font.size = Pt(9)

    heading("Headline numbers", 13)
    for line in [
        f"{c['eligible']} eligible — clear the 90/30/60 credit floors (incl. EC exemptions) and all named-course requirements.",
        f"{c['critical']} would be wrongly passed by a credit-only check — meet credit thresholds but missing a named required course.",
        f"{c['near']} near — named requirements satisfied, only 1–2 courses short on credits.",
        f"Remaining {c['total'] - c['eligible'] - c['critical'] - c['near']} are further from completion (see spreadsheet).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ---- Eligible table ----
    heading("Eligible to graduate — with class of degree (weighted GPA primary)", 13)
    expl = doc.add_paragraph()
    _eras_txt = "; ".join(f"{lab} — {desc}" for lab, desc in all_band_eras())
    er = expl.add_run("Class of degree on the WEIGHTED GPA per Handbook regulation (Level I weighted zero; "
                      "Levels II & III, passed courses). Each student is classified on the bands in force in "
                      "their year of entry: " + _eras_txt +
                      ". 'Wt GPA' is this figure; '(Printed)' is the unweighted BOE "
                      "Degree GPA. Rows marked * are where the two give a different class — the weighted GPA "
                      "governs; confirm with the registry.")
    er.italic = True
    er.font.size = Pt(9)

    cols = ["Student", "Page", "Programme", "Entry", "Wt GPA", "(Printed)", "Class of Degree"]

    # Group eligible students by department, then programme (alphabetical within).
    from .department_map import department_of, DEPARTMENT_ORDER
    from collections import defaultdict
    band_rank = {"First Class Honours": 0, "Upper Second Class Honours": 1,
                 "Lower Second Class Honours": 2, "Pass": 3}
    grouped = defaultdict(lambda: defaultdict(list))
    for r in summary["eligible"]:
        grouped[department_of(r["programme"])][r["programme"]].append(r)

    dept_order = DEPARTMENT_ORDER + [d for d in grouped if d not in DEPARTMENT_ORDER]
    for dept in dept_order:
        if dept not in grouped:
            continue
        # Department subheading
        dh = doc.add_paragraph()
        dr = dh.add_run(dept)
        dr.bold = True
        dr.font.size = Pt(12)
        dr.font.color.rgb = navy
        for programme in sorted(grouped[dept]):
            rows = sorted(grouped[dept][programme],
                          key=lambda x: (x.get("page") if x.get("page") is not None else 10**9))
            ph = doc.add_paragraph()
            pr = ph.add_run(programme + f"  ({len(rows)})")
            pr.bold = True
            pr.italic = True
            pr.font.size = Pt(10)
            t = doc.add_table(rows=1, cols=len(cols))
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.style = "Table Grid"
            for i, h in enumerate(cols):
                cell = t.rows[0].cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
                shade(cell, NAVY)
            for r in rows:
                cells = t.add_row().cells
                star = " *" if r.get("class_differs") else ""
                wg = r.get("weighted_gpa")
                vals = [r["name"], str(r.get("page", "")), r["programme"], r["entry_year"],
                        f"{wg:.2f}" if wg is not None else "—",
                        (f"{r['gpa']:.2f}" if r.get("gpa") is not None else "—"),
                        (r.get("class_of_degree", "") + star)]
                for i, v in enumerate(vals):
                    cells[i].text = ""
                    run = cells[i].paragraphs[0].add_run(v)
                    run.font.size = Pt(9)
                    if i == 6:
                        run.bold = True
                        shade(cells[i], "FFF2CC" if r.get("class_differs") else "E2EFDA")

    diffs = [r for r in summary["eligible"] if r.get("class_differs")]
    if diffs:
        names = "; ".join(f"{r['name'].split(',')[0]} ({r['class_printed_gpa'].split()[0]} "
                          f"→ {r['class_of_degree'].split()[0]} {r['class_of_degree'].split()[1]})"
                          for r in diffs)
        p = doc.add_paragraph()
        pr = p.add_run(f"* Weighted GPA gives a different class than the printed Degree GPA for: {names}. "
                       "These should be confirmed with the registry.")
        pr.italic = True
        pr.font.size = Pt(9)

    # ---- Critical table ----
    heading("Critical — credit totals met, but a named required course is missing", 13)
    p = doc.add_paragraph()
    pr = p.add_run("A purely numeric check would wrongly clear these. None can be recommended until the named "
                   "course is completed.")
    pr.italic = True
    pr.font.size = Pt(9)
    cols = ["Student", "Page", "Programme", "Entry", "Credits", "Missing required course(s)"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    for i, h in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shade(cell, NAVY)
    for r in summary["critical"]:
        cells = t.add_row().cells
        miss = ", ".join(r.get("missing_level1", []) + r.get("missing_level23_named", []))
        vals = [r["name"], str(r.get("page", "")), r["programme"], r["entry_year"],
                str(r.get("total_credits_effective", "")), miss]
        for i, v in enumerate(vals):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(v)
            run.font.size = Pt(9)
            if i == 5:
                run.bold = True
                shade(cells[i], "FCE4D6")

    # ---- Near table ----
    heading("Near — named requirements met, only 1–2 courses short on credits", 13)
    cols = ["Student", "Page", "Programme", "Entry", "Credits", "Courses short (est.)"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    for i, h in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shade(cell, NAVY)
    for r in summary["near"]:
        cells = t.add_row().cells
        vals = [r["name"], str(r.get("page", "")), r["programme"], r["entry_year"],
                str(r.get("total_credits_effective", "")), f"~{r['n']}"]
        for i, v in enumerate(vals):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(v)
            run.font.size = Pt(9)

    # ---- Method ----
    heading("Method & limits", 13)
    for line in [
        "Records are split per student, not per page (two students may share a physical page).",
        "Entry-year rule: each student is held to the programme structure from their year of entry.",
        "EC exemptions count toward the credit floors; the writing requirement (FOUN 1006/1008) is separate "
        "from the FOUN 1101/1201/Foreign-Language requirement.",
        "Class of degree uses the weighted GPA (Level I weighted zero; Levels II & III, passed courses).",
        "Not decided here: elective-slot counts, the Management 30/30 split, EX-exemption validity, "
        "rate-of-progress limits, BOE annotations, or substitutions/waivers. Final sign-off remains with the examiner.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def build_pdf(summary, path, docx_path=None):
    """
    Prefer converting the docx to PDF via LibreOffice (keeps identical layout).
    Fall back to a direct ReportLab PDF if LibreOffice isn't available.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if docx_path and soffice and os.path.exists(docx_path):
        outdir = os.path.dirname(path) or "."
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                check=True, capture_output=True, timeout=120,
            )
            produced = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if os.path.exists(produced):
                if produced != path:
                    shutil.move(produced, path)
                return path
        except Exception:
            pass
    return _build_pdf_reportlab(summary, path)


def _build_pdf_reportlab(summary, path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    navy = colors.HexColor("#1F4E78")
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=navy, fontSize=18)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=navy, fontSize=13)
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=9)
    ital = ParagraphStyle("ital", parent=styles["Normal"], fontSize=8, textColor=colors.grey)

    doc = SimpleDocTemplate(path, pagesize=letter,
                            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
                            topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    elems = []
    c = summary["counts"]
    elems.append(Paragraph("Level 3 Graduation Eligibility Review", h1))
    elems.append(Paragraph(f"Faculty of Social Sciences — {c['total']} Level 3 records assessed", body))
    elems.append(Spacer(1, 6))
    elems.append(Paragraph(
        f"<b>{c['eligible']}</b> eligible &nbsp; | &nbsp; <b>{c['critical']}</b> credit-OK but missing a named "
        f"course &nbsp; | &nbsp; <b>{c['near']}</b> near (1–2 short)", body))
    elems.append(Spacer(1, 10))

    def make_table(data, header_bg=navy, col_widths=None):
        t = Table(data, colWidths=col_widths, repeatRows=1)
        ts = [
            ("BACKGROUND", (0, 0), (-1, 0), header_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F7F7")]),
        ]
        t.setStyle(TableStyle(ts))
        return t

    # Eligible
    elems.append(Paragraph("Eligible — with class of degree (weighted GPA)", h2))
    data = [["Student", "Pg", "Programme", "Entry", "WtGPA", "(Print)", "Class of Degree"]]
    for r in summary["eligible"]:
        wg = r.get("weighted_gpa")
        star = " *" if r.get("class_differs") else ""
        data.append([r["name"], str(r.get("page", "")), r["programme"], r["entry_year"],
                     f"{wg:.2f}" if wg is not None else "—", f"{r['gpa']:.2f}",
                     r.get("class_of_degree", "") + star])
    elems.append(make_table(data, col_widths=[1.5 * inch, 0.35 * inch, 1.3 * inch,
                                              0.6 * inch, 0.5 * inch, 0.5 * inch, 1.85 * inch]))
    elems.append(Spacer(1, 10))

    # Critical
    elems.append(Paragraph("Critical — credit totals met, but a named course missing", h2))
    data = [["Student", "Pg", "Programme", "Entry", "Cr", "Missing required course(s)"]]
    for r in summary["critical"]:
        miss = ", ".join(r.get("missing_level1", []) + r.get("missing_level23_named", []))
        data.append([r["name"], str(r.get("page", "")), r["programme"], r["entry_year"],
                     str(r.get("total_credits_effective", "")), miss])
    elems.append(make_table(data, col_widths=[1.5 * inch, 0.35 * inch, 1.2 * inch,
                                              0.6 * inch, 0.35 * inch, 2.6 * inch]))
    elems.append(Spacer(1, 10))

    # Near
    elems.append(Paragraph("Near — named requirements met, 1–2 courses short", h2))
    data = [["Student", "Pg", "Programme", "Entry", "Cr", "Short"]]
    for r in summary["near"]:
        data.append([r["name"], str(r.get("page", "")), r["programme"], r["entry_year"],
                     str(r.get("total_credits_effective", "")), f"~{r['n']}"])
    elems.append(make_table(data, col_widths=[1.7 * inch, 0.4 * inch, 1.5 * inch,
                                              0.7 * inch, 0.5 * inch, 0.7 * inch]))
    elems.append(Spacer(1, 10))
    elems.append(Paragraph(
        "Class of degree uses the weighted GPA (Level I weighted zero; Levels II & III, passed courses). "
        "Rows marked * differ from the printed Degree GPA and should be confirmed with the registry. "
        "Credit floors include EC exemptions. Final sign-off remains with the examiner.", ital))

    doc.build(elems)
    return path
