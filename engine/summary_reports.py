"""
Builders for the summary reports:
  build_summary_docx(results, students, path)
  build_summary_xlsx(results, students, path)

Both contain:
  1. First Class Honours students, by department & programme (name, ID, GPA)
  2. Degree-band counts by programme (eligible, last 3 academic years)
  3. Semester counts by programme (eligible, last 3 academic years)
  4. First Class counts by department
  5. Current-year eligible counts by semester x degree band
"""
from . import analytics as A
from .department_map import DEPARTMENT_ORDER

NAVY = "1F4E78"
BANDS = A.BANDS_ORDER
SEMS = ["Semester 1", "Semester 2", "Summer"]


def _prep(results, students):
    A.enrich(results, students)
    cy = A.current_academic_year(students)
    years = A.last_n_years(cy, 3)
    return cy, years


# ===========================================================================
# DOCX
# ===========================================================================
def build_summary_docx(results, students, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    navy = RGBColor(0x1F, 0x4E, 0x78)
    cy, years = _prep(results, students)

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
        tcPr.append(sh)

    def H(text, size=15, after=4):
        p = doc.add_paragraph(); p.space_after = Pt(after)
        r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = navy
        return p

    def table(headers, rows, widths=None, total_row=None):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = ""
            run = c.paragraphs[0].add_run(str(h)); run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(9)
            shade(c, NAVY)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run("" if v is None else str(v))
                run.font.size = Pt(9)
        if total_row:
            cells = t.add_row().cells
            for i, v in enumerate(total_row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run("" if v is None else str(v))
                run.bold = True; run.font.size = Pt(9)
                shade(cells[i], "DDEBF7")
        return t

    # Title
    title = doc.add_paragraph()
    tr = title.add_run("Level 3 Graduation — Summary Report")
    tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = navy
    sub = doc.add_paragraph()
    sr = sub.add_run(f"Faculty of Social Sciences · Current academic year {cy} · "
                     f"Three-year window: {years[0]}–{years[-1]}")
    sr.italic = True; sr.font.size = Pt(10)
    note = doc.add_paragraph()
    nr = note.add_run("All counts below are ELIGIBLE students only. Class of degree is on the "
                      "weighted GPA. Graduating semester/year is derived from each student's "
                      "latest course term.")
    nr.italic = True; nr.font.size = Pt(8.5)

    # 1. First Class by department & programme
    H("1.  First Class Honours — students by department and programme", 14)
    fc = A.first_class_by_programme(results)
    dept_order = DEPARTMENT_ORDER + [d for d in fc if d not in DEPARTMENT_ORDER]
    any_fc = False
    for dept in dept_order:
        if dept not in fc:
            continue
        any_fc = True
        dh = doc.add_paragraph(); dr = dh.add_run(dept)
        dr.bold = True; dr.font.size = Pt(12); dr.font.color.rgb = navy
        for prog in sorted(fc[dept]):
            students_fc = sorted(fc[dept][prog], key=lambda x: -(x["gpa"] or 0))
            ph = doc.add_paragraph(); pr = ph.add_run(f"{prog}  ({len(students_fc)})")
            pr.bold = True; pr.italic = True; pr.font.size = Pt(10)
            rows = [[x["name"], x["id"],
                     f"{x['gpa']:.2f}" if x["gpa"] is not None else "—",
                     f"{x['printed_gpa']:.2f}" if x["printed_gpa"] is not None else "—",
                     x["page"]] for x in students_fc]
            table(["Student", "Student ID", "Wt GPA", "Printed GPA", "Page"], rows)
    if not any_fc:
        doc.add_paragraph("No First Class Honours students in this batch.")

    # 2. Degree bands by programme (3-year window)
    doc.add_page_break()
    H(f"2.  Degree-band counts by programme  (eligible, {years[0]}–{years[-1]})", 14)
    bbp = A.bands_by_programme(results, years)
    for dept in DEPARTMENT_ORDER:
        if dept not in bbp:
            continue
        dh = doc.add_paragraph(); dr = dh.add_run(dept)
        dr.bold = True; dr.font.size = Pt(12); dr.font.color.rgb = navy
        rows = []
        dept_tot = [0, 0, 0, 0]
        for prog in sorted(bbp[dept]):
            counts = bbp[dept][prog]
            vals = [counts.get(b, 0) for b in BANDS]
            dept_tot = [a + b for a, b in zip(dept_tot, vals)]
            rows.append([prog] + vals + [sum(vals)])
        table(["Programme", "First", "Upper 2nd", "Lower 2nd", "Pass", "Total"], rows,
              total_row=["Department total"] + dept_tot + [sum(dept_tot)])

    # 3. Semester by programme (3-year window)
    doc.add_page_break()
    H(f"3.  Graduating-semester counts by programme  (eligible, {years[0]}–{years[-1]})", 14)
    sbp = A.semester_by_programme(results, years)
    for dept in DEPARTMENT_ORDER:
        if dept not in sbp:
            continue
        dh = doc.add_paragraph(); dr = dh.add_run(dept)
        dr.bold = True; dr.font.size = Pt(12); dr.font.color.rgb = navy
        rows = []
        dept_tot = [0, 0, 0]
        for prog in sorted(sbp[dept]):
            counts = sbp[dept][prog]
            vals = [counts.get(sm, 0) for sm in SEMS]
            dept_tot = [a + b for a, b in zip(dept_tot, vals)]
            rows.append([prog] + vals + [sum(vals)])
        table(["Programme", "Semester 1", "Semester 2", "Summer", "Total"], rows,
              total_row=["Department total"] + dept_tot + [sum(dept_tot)])

    # 4. First Class by department
    doc.add_page_break()
    H("4.  First Class Honours — count by department", 14)
    fcd = A.first_class_by_department(results)
    rows = [[d, fcd.get(d, 0)] for d in DEPARTMENT_ORDER]
    table(["Department", "First Class Honours"], rows,
          total_row=["Total", sum(fcd.values())])

    # 5. Current-year eligible by semester x band
    H(f"5.  Eligible students in {cy} — by semester and degree band", 14)
    cyb = A.current_year_semester_bands(results, cy)
    rows = []
    col_tot = [0, 0, 0, 0]
    for sm in SEMS:
        vals = [cyb.get(sm, {}).get(b, 0) for b in BANDS]
        col_tot = [a + b for a, b in zip(col_tot, vals)]
        rows.append([sm] + vals + [sum(vals)])
    table(["Semester", "First", "Upper 2nd", "Lower 2nd", "Pass", "Total"], rows,
          total_row=["Total"] + col_tot + [sum(col_tot)])

    doc.save(path)
    return path


# ===========================================================================
# XLSX
# ===========================================================================
def build_summary_xlsx(results, students, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    cy, years = _prep(results, students)

    wb = Workbook()
    header_font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    header_fill = PatternFill("solid", start_color=NAVY)
    title_font = Font(bold=True, size=13, name="Arial", color="1F4E78")
    bold = Font(bold=True, name="Arial", size=10)
    normal = Font(name="Arial", size=10)
    totfill = PatternFill("solid", start_color="DDEBF7")
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def style_header_row(ws, row, ncols):
        for c in range(1, ncols + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = header_font; cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
            cell.border = border

    def widths(ws, ws_widths):
        for i, w in enumerate(ws_widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w

    # --- Sheet 1: First Class list ---
    ws = wb.active; ws.title = "First Class List"
    ws["A1"] = f"First Class Honours — by department & programme ({cy})"
    ws["A1"].font = title_font
    ws.append([]); 
    hdr = ["Department", "Programme", "Student", "Student ID", "Wt GPA", "Printed GPA", "Page"]
    ws.append(hdr); style_header_row(ws, ws.max_row, len(hdr))
    fc = A.first_class_by_programme(results)
    for dept in DEPARTMENT_ORDER:
        if dept not in fc:
            continue
        for prog in sorted(fc[dept]):
            for x in sorted(fc[dept][prog], key=lambda z: -(z["gpa"] or 0)):
                ws.append([dept, prog, x["name"], x["id"],
                           round(x["gpa"], 2) if x["gpa"] is not None else None,
                           round(x["printed_gpa"], 2) if x["printed_gpa"] is not None else None,
                           x["page"]])
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for c in row:
            c.font = normal; c.border = border
    widths(ws, [22, 26, 28, 12, 9, 11, 7])
    ws.freeze_panes = "A4"

    # --- Sheet 2: Bands by programme ---
    ws = wb.create_sheet("Bands by Programme")
    ws["A1"] = f"Degree-band counts by programme (eligible, {years[0]}-{years[-1]})"
    ws["A1"].font = title_font
    ws.append([])
    hdr = ["Department", "Programme", "First", "Upper 2nd", "Lower 2nd", "Pass", "Total"]
    ws.append(hdr); style_header_row(ws, ws.max_row, len(hdr))
    bbp = A.bands_by_programme(results, years)
    grand = [0, 0, 0, 0]
    for dept in DEPARTMENT_ORDER:
        if dept not in bbp:
            continue
        dtot = [0, 0, 0, 0]
        for prog in sorted(bbp[dept]):
            vals = [bbp[dept][prog].get(b, 0) for b in BANDS]
            dtot = [a + b for a, b in zip(dtot, vals)]
            ws.append([dept, prog] + vals + [sum(vals)])
        r = ws.max_row + 1
        ws.append([dept + " total", ""] + dtot + [sum(dtot)])
        for c in ws[ws.max_row]:
            c.font = bold; c.fill = totfill
        grand = [a + b for a, b in zip(grand, dtot)]
    ws.append(["FACULTY TOTAL", ""] + grand + [sum(grand)])
    for c in ws[ws.max_row]:
        c.font = bold; c.fill = totfill
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for c in row:
            c.border = border
            if not c.font.bold:
                c.font = normal
    widths(ws, [22, 28, 8, 10, 10, 8, 8])
    ws.freeze_panes = "A4"

    # --- Sheet 3: Semester by programme ---
    ws = wb.create_sheet("Semester by Programme")
    ws["A1"] = f"Graduating-semester counts by programme (eligible, {years[0]}-{years[-1]})"
    ws["A1"].font = title_font
    ws.append([])
    hdr = ["Department", "Programme", "Semester 1", "Semester 2", "Summer", "Total"]
    ws.append(hdr); style_header_row(ws, ws.max_row, len(hdr))
    sbp = A.semester_by_programme(results, years)
    grand = [0, 0, 0]
    for dept in DEPARTMENT_ORDER:
        if dept not in sbp:
            continue
        dtot = [0, 0, 0]
        for prog in sorted(sbp[dept]):
            vals = [sbp[dept][prog].get(sm, 0) for sm in SEMS]
            dtot = [a + b for a, b in zip(dtot, vals)]
            ws.append([dept, prog] + vals + [sum(vals)])
        ws.append([dept + " total", ""] + dtot + [sum(dtot)])
        for c in ws[ws.max_row]:
            c.font = bold; c.fill = totfill
        grand = [a + b for a, b in zip(grand, dtot)]
    ws.append(["FACULTY TOTAL", ""] + grand + [sum(grand)])
    for c in ws[ws.max_row]:
        c.font = bold; c.fill = totfill
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for c in row:
            c.border = border
            if not c.font.bold:
                c.font = normal
    widths(ws, [22, 28, 11, 11, 9, 8])
    ws.freeze_panes = "A4"

    # --- Sheet 4: Department roll-ups ---
    ws = wb.create_sheet("Department Summary")
    ws["A1"] = "First Class Honours — by department"
    ws["A1"].font = title_font
    ws.append([]); ws.append(["Department", "First Class Honours"])
    style_header_row(ws, ws.max_row, 2)
    fcd = A.first_class_by_department(results)
    for d in DEPARTMENT_ORDER:
        ws.append([d, fcd.get(d, 0)])
    ws.append(["Total", sum(fcd.values())])
    for c in ws[ws.max_row]:
        c.font = bold; c.fill = totfill
    start2 = ws.max_row + 3
    ws.cell(row=start2, column=1, value=f"Eligible in {cy} — by semester & degree band").font = title_font
    hdr_row = start2 + 1
    for i, h in enumerate(["Semester", "First", "Upper 2nd", "Lower 2nd", "Pass", "Total"], start=1):
        ws.cell(row=hdr_row, column=i, value=h)
    style_header_row(ws, hdr_row, 6)
    cyb = A.current_year_semester_bands(results, cy)
    coltot = [0, 0, 0, 0]
    rr = hdr_row + 1
    for sm in SEMS:
        vals = [cyb.get(sm, {}).get(b, 0) for b in BANDS]
        coltot = [a + b for a, b in zip(coltot, vals)]
        for i, v in enumerate([sm] + vals + [sum(vals)], start=1):
            ws.cell(row=rr, column=i, value=v)
        rr += 1
    for i, v in enumerate(["Total"] + coltot + [sum(coltot)], start=1):
        cell = ws.cell(row=rr, column=i, value=v); cell.font = bold; cell.fill = totfill
    for row in ws.iter_rows(min_row=3, max_row=ws.max_row):
        for c in row:
            if c.value is not None:
                c.border = border
                if not c.font.bold:
                    c.font = normal
    widths(ws, [22, 10, 11, 11, 8, 8])

    wb.save(path)
    return path


# ===========================================================================
# Cross-period (master) summary — built from stored, deduped master rows
# ===========================================================================
def build_master_summary_xlsx(master_rows, path):
    """Cross-period summary across all stored uploads (deduped to most-complete)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from . import analytics as A

    wb = Workbook()
    header_font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    header_fill = PatternFill("solid", start_color=NAVY)
    title_font = Font(bold=True, size=13, name="Arial", color="1F4E78")
    bold = Font(bold=True, name="Arial", size=10)
    normal = Font(name="Arial", size=10)
    totfill = PatternFill("solid", start_color="DDEBF7")
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def style_header(ws, row, n):
        for c in range(1, n + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = header_font; cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", wrap_text=True); cell.border = border

    def widths(ws, ws_w):
        for i, w in enumerate(ws_w, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w

    years = sorted({r.get("grad_year") for r in master_rows if r.get("grad_year")})

    # Sheet 1: bands by year
    ws = wb.active; ws.title = "Bands by Year"
    ws["A1"] = "Cross-period: degree bands by graduating year (deduped, eligible)"
    ws["A1"].font = title_font
    ws.append([]); hdr = ["Graduating year"] + BANDS + ["Total"]
    ws.append(hdr); style_header(ws, ws.max_row, len(hdr))
    bby = A.master_bands_by_year(master_rows)
    coltot = [0, 0, 0, 0]
    for y in years:
        vals = [bby.get(y, {}).get(b, 0) for b in BANDS]
        coltot = [a + b for a, b in zip(coltot, vals)]
        ws.append([y] + vals + [sum(vals)])
    ws.append(["Total"] + coltot + [sum(coltot)])
    for c in ws[ws.max_row]:
        c.font = bold; c.fill = totfill
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for c in row:
            c.border = border
            if not c.font.bold:
                c.font = normal
    widths(ws, [18, 10, 11, 11, 8, 8])

    # Sheet 2: eligible by year x department
    ws = wb.create_sheet("Eligible by Year & Dept")
    ws["A1"] = "Cross-period: eligible students by year and department (deduped)"
    ws["A1"].font = title_font
    ws.append([]); hdr = ["Graduating year"] + DEPARTMENT_ORDER + ["Total"]
    ws.append(hdr); style_header(ws, ws.max_row, len(hdr))
    eyd = A.master_eligible_by_year_dept(master_rows)
    coltot = [0] * len(DEPARTMENT_ORDER)
    for y in years:
        vals = [eyd.get(y, {}).get(d, 0) for d in DEPARTMENT_ORDER]
        coltot = [a + b for a, b in zip(coltot, vals)]
        ws.append([y] + vals + [sum(vals)])
    ws.append(["Total"] + coltot + [sum(coltot)])
    for c in ws[ws.max_row]:
        c.font = bold; c.fill = totfill
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for c in row:
            c.border = border
            if not c.font.bold:
                c.font = normal
    widths(ws, [18] + [20] * len(DEPARTMENT_ORDER) + [8])

    # Sheet 3: First Class by department (all-time, deduped)
    ws = wb.create_sheet("First Class by Dept")
    ws["A1"] = "Cross-period: First Class Honours by department (deduped)"
    ws["A1"].font = title_font
    ws.append([]); ws.append(["Department", "First Class Honours"])
    style_header(ws, ws.max_row, 2)
    fcd = A.master_first_class_by_department(master_rows)
    for d in DEPARTMENT_ORDER:
        ws.append([d, fcd.get(d, 0)])
    ws.append(["Total", sum(fcd.values())])
    for c in ws[ws.max_row]:
        c.font = bold; c.fill = totfill
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for c in row:
            c.border = border
            if not c.font.bold:
                c.font = normal
    widths(ws, [22, 18])

    wb.save(path)
    return path
